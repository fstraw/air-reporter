# -*- coding: utf-8 -*-
"""
Created on Tue Jun 02 15:43:00 2015

@author: bbatt
"""
import os

import docx
from win32com import client

document = docx.Document(r'C:\Users\bbatt\Dropbox\!Python\air-reporter\auxfiles\appendix_workable.docx')
pi = ''
pinum = 'PI No. {}'.format(pi)
appendices = (('Appendix A', 'Traffic'),
            ('Appendix B', 'MOVES'),
            ('Appendix C', 'CO Model'),
            ('Appendix D', 'PM Correspondence'))

ws = r"C:\Users\bbatt\Dropbox\!Python\air-reporter\airreporter"

def ConvertToPDF(doc):
    try:
        word = client.DispatchEx("Word.Application")
        new_name = doc.replace(".docx", r".pdf")
        worddoc = word.Documents.Open(doc)
        worddoc.SaveAs(new_name, FileFormat = 17)
        worddoc.Close()
        return new_name
    except Exception, e:
        return e
    finally:
        word.Quit()

#create word docx of splitters
def add_appendix(appnum, appname, pinum):
    document.add_heading(appnum, 1)
    document.add_heading(appname, 2)
    document.add_paragraph(pinum)
    document.add_page_break()

#create pdf of splitters
def create_splitters(ws, appendices, pinum):
    for appendix in appendices:
        add_appendix(appendix[0], appendix[1], pinum)
    docname = os.path.join(ws, 'appendix.docx')
    document.save(docname)
    result = ConvertToPDF(docname)
    os.remove(docname)
    return result

if __name__ == '__main__':
    create_splitters(ws, appendices, pinum)
