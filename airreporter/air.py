import docx, time, sys, os, ConfigParser
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from math import ceil

"""To Do
Add TIP region picker
TIP use cases in document (ozone section in particular)
SIP reference in Executive summary
Change Ozone discussion to specify region
"""
#Specify template document for styles and word document with project description
# Find and gather settings from the ini file

localpath = r"C:\Users\bbatt\Dropbox\!Python\air-reporter\auxfiles"
settingsFile = os.path.join(localpath, "projectinfo.ini")

if os.path.isfile(settingsFile):
    config = ConfigParser.ConfigParser()
    config.read(settingsFile)
else:
    print "INI file not found. Make sure a valid 'settings.ini' file exists in the same directory as this script."
    sys.exit()
#DOCX file containing styles
document = docx.Document('../auxfiles/report_template.docx')
#Project description filler
project_desc = 'Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed sed tristique felis. Nullam scelerisque ligula eu magna elementum interdum. In sed consequat turpis, porttitor mollis leo. Suspendisse blandit tincidunt ante, aliquam semper lacus volutpat sit amet. Maecenas scelerisque maximus metus eu mattis. Curabitur sit amet nibh bibendum, vehicula ante a, porta lacus. Vestibulum elementum cursus ligula, suscipit consequat augue luctus quis. Nulla facilisi. Nam quis sem ex. Maecenas accumsan ligula ut lorem dapibus, in sagittis sapien dignissim. In eget fermentum nisl. Sed pretium nibh lectus, ut accumsan dolor volutpat nec. Ut sed est non augue vulputate vehicula ac ut felis. Sed accumsan tristique dolor nec maximus. Ut bibendum massa sit amet rhoncus mollis.'
# Report values
preparedBy = config.get('REPORT_INFO', 'PREPAREDBY')
qcBy = config.get('REPORT_INFO', 'QCBY')
strReportType = config.get('REPORT_INFO', 'REPORTTYPE')
strProjectName = config.get('REPORT_INFO', 'PROJECTNAME')
strGDOTProj = config.get('REPORT_INFO', 'GDOTPROJ')
strGDOTPI = config.get('REPORT_INFO', 'GDOTPI')
strMPO = config.get('REPORT_INFO', 'MPO')
strTIP = config.get('REPORT_INFO', 'TIP')
strCounty = config.get('REPORT_INFO', 'COUNTY')
strConcDate = config.get('REPORT_INFO', 'CONCDATE')
strProjSum = config.get('REPORT_INFO', 'PROJSUM')
boolCOReq = bool(config.get('REPORT_INFO', 'COREQ'))
boolMSAT = bool(config.get('REPORT_INFO', 'MEANINGFULMSAT'))
boolOzone = bool(config.get('REPORT_INFO', 'OZONEATTAINMENT'))
boolPM = boolOzone
strPMDet = bool(config.get('REPORT_INFO', 'LODNEEDED'))
strProjDesc = config.get('REPORT_INFO', 'PROJECTDESC')
##MSAT section...[BLANK] will have the effect of moving some traffic closer to nearby homes and businesses;
strMSATDesc = config.get('REPORT_INFO', 'MSATDESC')

#MSAT section...The localized increases in MSAT concentrations would likely be most pronounced along [BLANK]
strMSATConc = config.get('REPORT_INFO', 'MSATCONC')

#CO Maximums
strCOWorkspace = config.get('CO_REPORT', 'COWORKSPACE')
strIntersection = config.get('CO_REPORT', 'INTERSECTION')
strCOBuild = config.get('CO_REPORT', 'COBUILD')
strCONoBuild = config.get('CO_REPORT', 'CONOBUILD')
strExistingYear = config.get('CO_REPORT', 'EXISTINGYEAR')
strDesignYear = config.get('CO_REPORT', 'DESIGNYEAR')
strMaxReceptorBuild = config.get('CO_REPORT', 'MAXRECEPTORBUILD')
strMaxReceptorNoBuild = config.get('CO_REPORT', 'MAXRECEPTORNOBUILD')
strReceptorQuadrant = config.get('CO_REPORT', 'RECEPTORQUADRANTS')

#Number of intersections operating at LOS C,D, or F
strIntersectionNo = config.get('CO_REPORT', 'LOSINTERSECTIONS')
strRoadType = config.get('CO_REPORT', 'ROADTYPE')
strSpeedLimitRange = config.get('CO_REPORT', 'SPEEDLIMITRANGE')
#Select Region - (7-county region, 13-county region, 45 county region, urban other, or rural other)
strRegion = config.get('CO_REPORT', 'REGION')
strAuto = config.get('CO_REPORT', 'AUTO')
strMedium = config.get('CO_REPORT', 'MEDIUM')
strHeavy = config.get('CO_REPORT', 'HEAVY')
strStabilityClass = config.get('CO_REPORT', 'STABILITYCLASS')
strSurfaceRoughness = config.get('CO_REPORT', 'SURFACEROUGHNESS')
#Background concentration - (1 - rural, 2 - suburban, 3 - urban)
strBackgroundConc = config.get('CO_REPORT', 'BACKGROUNDCONC')
strBuildRec = config.get('CO_REPORT', 'RECEPTORSBUILD')
strNoBuildRec = config.get('CO_REPORT', 'RECEPTORSNOBUILD')

#MSAT
strRoadway = config.get('MSAT_REPORT', 'ROAD')
dblLength = float(config.get('MSAT_REPORT', 'LENGTH'))
intExistingADT = int(config.get('MSAT_REPORT', 'EXISTINGADT'))
intNoBuildADT = int(config.get('MSAT_REPORT', 'NOBUILDADT'))
intBuildADT = int(config.get('MSAT_REPORT', 'BUILDADT'))

#existingout = os.path.join(workspace,"Existing.out")
#nobuildout = os.path.join(workspace,"NoBuild.out")
#buildout =  os.path.join(workspace,"Build.out")
#outputs = (existingout,nobuildout,buildout)
#document = docx.Document()

red = RGBColor(0xff, 0, 0)
def maxreceptorcount(tploutputs):
    x = 0
    for item in tploutputs:
        if len(item) > x:
            x = len(item)
    return x

def parsecoresults(workspace, backgroundppm):
    existingout = os.path.join(workspace,"Existing.out")
    nobuildout = os.path.join(workspace,"NoBuild.out")
    buildout =  os.path.join(workspace,"Build.out")
    outputs = (existingout,nobuildout,buildout)    
    alloutputs = []
    for output in outputs:            
        f = open(output, "rb")
        with f:
            linelist = [i for i in f.readlines() if i.startswith(" MAX")]
            cooutputlist = [i for i in linelist[0].split() if i != '*' and i != "MAX"]
            convertedlist = [float(i) + backgroundppm for i in cooutputlist]
        alloutputs.append(convertedlist)
    return alloutputs

def insertcographics(doc, workspace):
    document = doc
    table = document.add_table(rows=1, cols=2)
    row_cells = table.rows[0].cells
    existingp = row_cells[0].paragraphs[0]
    existingp.style = 'TblCentered'
    existingr = existingp.add_run()
    if os.path.exists(os.path.join(workspace,'existing.png')):
        existingr.add_picture(os.path.join(workspace,'existing.png'), width = Inches(3.32))
    else:
        print "Could not find existing CO graphic in {}".format(workspace)
    existingp.add_run("Figure 2. Existing/NoBuild")
    buildp = row_cells[1].paragraphs[0]
    buildp.style = 'TblCentered'
    buildr = buildp.add_run()
    if os.path.exists(os.path.join(workspace,'build.png')):
        buildr.add_picture(os.path.join(workspace,'build.png'), width = Inches(3.32), height = Inches(3.32))
    else:
        print "Could not find build CO graphic in {}".format(workspace)
    buildp.add_run("Figure 3. Build")

def createcotable(doc, workspace, backgroundppm):
    document = doc
    table = document.add_table(rows=1, cols=4)
    table.style = 'LightShading'
    hdrcells = table.rows[0].cells
    paragraph = hdrcells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    paragraph.add_run("Receptors")
    paragraph = hdrcells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    paragraph.add_run("Existing")
    paragraph = hdrcells[2].paragraphs[0]
    paragraph.style = 'TblCentered'
    paragraph.add_run("No Build")
    paragraph = hdrcells[3].paragraphs[0]
    paragraph.style = 'TblCentered'
    paragraph.add_run("Build")
    if not os.path.exists(workspace):
        print "CO Directory {} does not exist!".format(workspace)
        return False
    outputsources = parsecoresults(workspace, backgroundppm)
    requiredtblrows = maxreceptorcount(outputsources)
    for i in range(requiredtblrows):
        row_cells = table.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        paragraph.style = 'TblCentered'
        paragraph.add_run("Receptor " + str(i+1))
        paragraph = row_cells[1].paragraphs[0]
        paragraph.style = 'TblCentered'
        paragraph.add_run(str(outputsources[0][i]))
        paragraph = row_cells[2].paragraphs[0]
        paragraph.style = 'TblCentered'
        paragraph.add_run(str(outputsources[1][i]))
        paragraph = row_cells[3].paragraphs[0]
        paragraph.style = 'TblCentered'
        paragraph.add_run(str(outputsources[2][i]))
    return table

def msattable(document, road, length, existing, nobuild, build):
	table = document.add_table(rows=3, cols=6)
	table.style = 'LightShading'
	hdrs = ["Roadway", "Roadway Length (Miles)", "", "Existing", "No Build", "Build" ]
	for i in range(len(hdrs)):
		table.rows[0].cells[i].text = hdrs[i]
	rowindex = 1
	roadcell = table.cell(rowindex,0).merge(table.cell(2,0))
	lengthcell = table.cell(rowindex, 1).merge(table.cell(2, 1))
	trafficlabelcell = table.cell(rowindex, 2)	
	existingcell = table.cell(rowindex,3)
	nobuildcell = table.cell(rowindex,4)
	buildcell = table.cell(rowindex,5)
	vmtlabelcell = table.cell(rowindex+ 1, 2)	
	existingvmtcell = table.cell(rowindex + 1, 3)
	nobuildvmtcell = table.cell(rowindex + 1, 4)
	buildvmtcell = table.cell(rowindex + 1, 5)
	roadcell.paragraphs[0].add_run(road)
	lengthcell.text = str(length)
	varlist = [roadcell, lengthcell, trafficlabelcell,
			existingcell, nobuildcell, buildcell, vmtlabelcell,
			existingvmtcell, nobuildvmtcell, buildvmtcell]
	for var in varlist:
		var.paragraphs[0].style = 'TblCentered'
	trafficlabelcell.paragraphs[0].add_run("ADT")
	existingcell.paragraphs[0].add_run("{:,}".format(existing))
	nobuildcell.paragraphs[0].add_run("{:,}".format(nobuild))
	buildcell.paragraphs[0].add_run("{:,}".format(build))
	vmtlabelcell.paragraphs[0].add_run("VMT")
	existingvmtcell.paragraphs[0].add_run("{:,}".format(ceil(existing * length)).rstrip(".").rstrip("0").rstrip("."))
	nobuildvmtcell.paragraphs[0].add_run("{:,}".format(ceil(nobuild * length)).rstrip(".").rstrip("0").rstrip("."))
	buildvmtcell.paragraphs[0].add_run("{:,}".format(ceil(build * length)).rstrip(".").rstrip("0").rstrip("."))
	return table

def executivesummary(document):
    p = document.add_paragraph()
    p.style = "Subtitle"
    p.add_run('Executive Summary').bold = True
    
    if strGDOTProj:
        p = document.add_paragraph()
        p.style = "Subtitle"
        #add if/then logic for multiple numbers
        p.add_run(strGDOTProj).bold = True
    
    p = document.add_paragraph()
    p.style = "Subtitle"
    p.add_run(strCounty + ", ").bold = True
    p.add_run(strGDOTPI).bold = True
    
    p = document.add_paragraph()
    p.style = "Subtitle"
    p.add_run(strProjectName).bold = True
    
    p = document.add_paragraph()
    p.style = "Subtitle"
    p.add_run(time.strftime("%B") + " " + time.strftime("%Y")).bold = True
    
    #Executive summary sections
    p = document.add_paragraph()
    p.style = "NoSpacing"
    p.add_run('Project Description: ').bold = True
    p.add_run(project_desc).font.color.rgb = red
    
    p = document.add_paragraph()
    p.style = "NoSpacing"
    if boolOzone:
    	p.add_run('Ozone: MPO and TIP Number: ').bold = True
    	p.add_run("""This project is identified in the %s and FY 2014-2019 Transportation Improvement Program by reference number(s) %s.""" % (strMPO, strTIP))
    else:
    	p.add_run('Ozone: ').bold = True
    	p.add_run("""This project is in an attainment area.""")     
    p = document.add_paragraph()
    p.style = "NoSpacing"
    p.add_run('PM').bold = True
    p.add_run('2.5').font.subscript = True
    p.add_run(': ').bold = True
    p.bold = True
    if boolPM:
    	p.add_run("""This project has been evaluated by an interagency group consisting of Federal Highway Administration (FHWA), Environmental Protection Agency (EPA), Georgia Department of Natural Resources Environmental Protection Division (GA EPD), and the Metropolitan Planning Organization (MPO) and they agreed that this project does NOT appear to be a "Project of Concern" per the Transportation Conformity Rule and thus meets the statutory and regulatory requirements for PM""")
    	p.add_run('2.5').font.subscript = True
    	p.add_run(""" hotspots without a quantitative analysis on %s.""" % (strConcDate))
    else:
    	p.add_run("""This project is in an attainment area.""")
    p = document.add_paragraph()
    p.style = "NoSpacing"
    p.add_run('CO Modeling Assumptions: ').bold = True
    if boolCOReq:
    	p.add_run("""The highest 1-hour CO concentration of %s ppm in the %s design year is projected at Receptor %s, located in the %s quadrant of the studied intersection. This value is lower than the maximum allowable NAAQS for the one-hour level of 35 ppm and the eight-hour level of 9 ppm.""" % (strCOBuild, strDesignYear, strMaxReceptorBuild, strReceptorQuadrant))
    else:
    	p.add_run("""The project was evaluated for the potential to result in increased CO concentrations in the project area. Based on LOS estimates, it has been determined that this project would not increase traffic congestion, increase idle emissions, or CO concentrations.""")
    
    p = document.add_paragraph()
    p.style = "NoSpacing"
    p.add_run('MSAT: ').bold = True
    if boolMSAT:
    	p.add_run('The proposed project is classified as a project with low meaningful MSAT effects.')
    else:
    	p.add_run('The proposed project is classified as a project with no meaningful MSAT effects.')
    
    p = document.add_paragraph()
    p.style = "NoSpacing"
    p.add_run('Conclusion: ').bold = True
    p.add_run('This project was evaluated for its consistency with state and federal air quality goals, including CO, Ozone, PM')
    p.add_run('2.5').font.subscript = True
    p.add_run(', and MSATs as part of this assessment. Results indicated that the project is consistent with the SIP for the attainment of clean air quality in Georgia and is in compliance with both state and federal air quality standards.')
    
    table = document.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style="TableGrid"
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Prepared By: %s' % (preparedBy)
    hdr_cells[1].text = 'QC/QA: %s' % (qcBy)
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = ' ' * 50
    hdr_cells[1].text = ' ' * 50
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = '_' * 40
    hdr_cells[1].text = '_' * 40
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = 'Signature' + ' ' * 50 + 'Date'
    hdr_cells[1].text = 'Signature' + ' ' * 50 + 'Date'
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = ' ' * 50
    hdr_cells[1].text = ' ' * 50
    hdr_cells = table.rows[5].cells
    hdr_cells[0].text = '_' * 40
    hdr_cells[1].text = ' ' * 50
    hdr_cells = table.rows[6].cells
    hdr_cells[0].text = 'Approved By: GDOT' + ' ' * 35
    hdr_cells = table.rows[7].cells
    hdr_cells[0].text = 'Signature' + ' ' * 50 + 'Date'

# document.add_page_break()
def reportbody(document):
    document.add_section()
    document.add_heading(strReportType, 1)
    document.add_heading(strProjectName, 1)
    document.add_heading(strGDOTPI, 1)
    document.add_heading(time.strftime("%B") + " " + time.strftime("%Y"), 1)
    
    document.add_heading('Introduction', 2)
    p = document.add_paragraph()
    p.add_run('The 1990 Clean Air Act (CAA) amendments and guidelines, issued by the U.S. Environmental Protection Agency (EPA), set forth guidelines to be followed by agencies responsible for attainment of the National Ambient Air Quality Standards (NAAQS). The CAA section 176(c) requires that Federal transportation projects are consistent with state air quality goals, found in the State Implementation Plan (SIP). The process to ensure this consistency is called Transportation Conformity. Conformity to the SIP means that transportation activities will not cause new violations of the NAAQS, worsen existing violations of the standards, or delay timely attainment of the relevant standard. In complying with these guidelines, the Georgia Department of Transportation (GDOT) has completed an analysis on the effects of the proposed project on air quality.')
    
    document.add_heading('What is the Proposed Project?', 2)
    p = document.add_paragraph()
    p.add_run(strProjSum)
    
    p = document.add_paragraph()
    p.add_run('Details:').underline = True
    
    p = document.add_paragraph()
    p.add_run(project_desc).font.color.rgb = red
    
    document.add_heading('What Criteria Pollutants Are Studied?', 2)
    p = document.add_paragraph()
    p.add_run('The NAAQS have been established for air pollutants that have been identified by the EPA as being of concern nationwide (')
    p.add_run('see Table 1 on page 3').font.color.rgb = red
    p.add_run('). These air pollutants, referred to as criteria pollutants, are carbon monoxide (CO), lead (Pb), nitrogen dioxide (NO')
    p.add_run('2').font.subscript = True
    p.add_run('), particulate matter (PM')
    p.add_run('2.5').font.subscript = True
    p.add_run('), ozone (O')
    p.add_run('3').font.subscript = True
    p.add_run('), and sulfur dioxide (SO')
    p.add_run('2').font.subscript = True
    p.add_run("""). The sources of these pollutants, effects on human health and the nation's welfare, and occurrence in the atmosphere vary considerably. In addition to the criteria air pollutants for which there are NAAQS, the EPA also regulates Mobile Source Air Toxics (MSATs).""")

    p = document.add_paragraph()
    p.style = "FigureCaption"
    p.add_run('Table 1. Pollutant Descriptions')
    table = document.add_table(rows=7, cols=2)
    table.style = 'LightShading'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Pollutant')
    paragraph = hdr_cells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Description')
     
    hdr_cells = table.rows[1].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Carbon Monoxide')
    paragraph = hdr_cells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Carbon monoxide is a colorless, odorless, poisonous gas formed when carbon in fuels is not burned completely.')
     
    hdr_cells = table.rows[2].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Lead')
    paragraph = hdr_cells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Lead is a chemical element in the carbon group. Lead is a soft, malleable and heavy post-transition metal.')
     
    hdr_cells = table.rows[3].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Nitrogen Dioxide')
    paragraph = hdr_cells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Nitrogen dioxide belongs to a family of highly reactive gases called nitrogen oxides (NOx). A suffocating, brownish gas, nitrogen dioxide is a strong oxidizing agent that reacts in the air to form corrosive nitric acid, as well as toxic organic nitrates.')
     
    hdr_cells = table.rows[4].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Particulate Matter')
    paragraph = hdr_cells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Particulate matter is the term for solid or liquid particles found in the air.')
     
    hdr_cells = table.rows[5].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Ozone')
    paragraph = hdr_cells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Ground-level ozone (the primary constituent of smog.) Ozone is created by sunlight acting on NOx and VOC in the air.')
     
    hdr_cells = table.rows[6].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('Sulfur Dioxide')
    paragraph = hdr_cells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    run = paragraph.add_run('These gases are formed when fuel containing sulfur (mainly coal and oil) is burned, and during metal smelting and other industrial processes.')
    
    p = document.add_paragraph()
    p.style = "SpaceBefore"
    p.add_run("""The EPA's regulations on Pb are focused on smelter and battery plants, the major sources of Pb in the air. A regulatory ban on the introduction of gasoline produced with the use of any lead additive was issued as of January 1, 1996. In the United States, sulfur levels in fuel have been slowly ratcheted down. The EPA's """)
    p.add_run("""Gasoline Sulfur""").font.underline = True
    p.add_run(""" program in effect (2004-2007) reduced the sulfur content of gasoline by up to 90 percent from uncontrolled levels. In addition, the Gasoline Sulfur program aims to lower sulfur content of gasoline to 10 parts per million (ppm) averages by 2017. A 15 ppm sulfur specification, known as Ultra Low Sulfur Diesel (ULSD), was phased in for highway diesel fuel from 2006-2010. In addition, Georgia is in attainment for Pb and SO""")
    p.add_run("""2""").font.subscript = True
    p.add_run("""; therefore, no further documentation is required.""")
    
    p = document.add_paragraph()
    p.add_run("""Georgia is in attainment for NO""")
    p.add_run("""2""").font.subscript = True
    p.add_run(""" & """)
    p.add_run("""PM""")
    p.add_run("""10""").font.subscript = True
    p.add_run("""; therefore, no further documentation is required.""")    

    p = document.add_paragraph()
    p.add_run("""Due to their association with roadway transportation sources, O""")
    p.add_run("""3""").font.subscript = True
    p.add_run(""", CO""")
    p.add_run(""", PM""")
    p.add_run("""2.5""").font.subscript = True
    p.add_run(""", and MSATs are typically reviewed for potential effects on nearby receptors with respect to roadway projects.""")     

    table = document.add_table(rows=1, cols=2)
    table.style = 'PMTableOutline'
    row_cells = table.rows[0].cells
    p = row_cells[0].paragraphs[0]
    p.style = 'TblCentered'
    p.add_run("""Particulate Matter (PM)""").bold = True  
    row_cells[0].add_paragraph()
    p = row_cells[0].add_paragraph("""Coarse dust particles (PM""")    
    p.add_run("""10""").font.subscript = True    
    p.add_run(""") are 2.5 to 10 micrometers in diameter. Sources include crushing or grinding operations and dust stirred up by vehicles on roads.""")
    row_cells[0].add_paragraph()
    p = row_cells[0].add_paragraph("""Fine particles (PM""")
    p.add_run("""2.5""").font.subscript = True
    p.add_run(""") are 2.5 micrometers in diameter or smaller, and can only be seen with an electron microscope. Fine particles are produced from all types of combustion, including motor vehicles, power plants, residential wood burning, forest fires, agricultural burning, and some industrial processes.""")
    buildp = row_cells[1].paragraphs[0]
    buildp.style = 'TblCentered'
    buildr = buildp.add_run()
    buildr.add_picture('../auxfiles/pm.jpg', width = Inches(4), height = Inches(3))

    hd = document.add_heading('Is this Project in an Ozone (O', 2)
    hd.add_run('3').font.subscript = True
    hd.add_run(') Non-Attainment Area?')
    p = document.add_paragraph()
    
    #Fix TIP Variables
    if boolOzone:
    	p.add_run("""This project is in an area where the SIP contains transportation control measures. The CAA requires Transportation Plans and Transportation Improvement Programs (TIP) in areas not meeting the NAAQS to conform to the emissions budget of the SIP for air quality. """)
    	p.add_run("""The FY 2014-2019 TIP is the current adopted plan for the Atlanta area showing the region's highest transportation priorities. It was adopted by the Atlanta Regional Commission on March 26, 2014, with GRTA Board action on April 9, 2014, and received conformity determination by the US DOT on April 30, 2014.""").font.color.rgb = red
    	
    	p = document.add_paragraph()
    	p.add_run("""This project is identified in the %s and FY 2014-2019 TIP by reference number %s.""" % (strMPO, strTIP)).font.color.rgb = red
    
    	p = document.add_paragraph()
    	p.add_run('Inclusion in a conforming plan also serves as project level analysis for O')
    	p.add_run('3').font.subscript = True
    	p.add_run('; no further analysis of O')
    	p.add_run('3').font.subscript = True
    	p.add_run(' emissions is warranted.')
    
    else:
    	p.add_run('This project is located outside of the ozone non-attainment area. Therefore, no project level analysis for O')
    	p.add_run('3').font.subscript = True
    	p.add_run(' is required.')
    
    document.add_heading('How Will The Project Affect Carbon Monoxide (CO) Emissions?', 2)
    
    if not boolCOReq:
    	p = document.add_paragraph()
    	p.add_run('Georgia is in attainment for CO; however, CO is also a concern in areas where signalized intersections (due to idling vehicles) are operating at a Level-of-Service (LOS) D, E, or F in the project design year (20 year design horizon).')
    
    	p = document.add_paragraph()
    	p.add_run("""The LOS is a standard means of classifying traffic conditions associated with various traffic volume levels and traffic flow conditions. There are six levels of service at which a roadway can operate, represented by the letters "A" through "F." Each level is defined by a maximum value for the ratio of traffic volume (V) to facility capacity (C) """)
    	p.add_run("""(see Table 2. LOS, below""").font.color.rgb = red     
    	p.add_run("""). The LOS for signalized intersections is determined by calculating the average control delay per vehicle for the intersection, i.e., the average amount of time it takes a vehicle to get through the intersection.""")
    
    	p = document.add_paragraph()
    	p.style = "FigureCaption"
    	p.add_run('Table 2. Level of Service (LOS)')
    	table = document.add_table(rows=7, cols=2)
    	table.style = 'LightShading'
    	table.alignment = WD_TABLE_ALIGNMENT.CENTER
    	hdr_cells = table.rows[0].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Level of Service')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Definition')
     
    	hdr_cells = table.rows[1].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('A')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('volume is well below capacity and traffic is flowing freely')
     
    	hdr_cells = table.rows[2].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('B')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('volume is steady, the presence of other vehicles begins to be noticeable')
     
    	hdr_cells = table.rows[3].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('C')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('steady traffic flow, speeds and maneuverability are more closely controlled by traffic volumes')
     
    	hdr_cells = table.rows[4].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('D')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('approaching an unsteady flow in which speed and maneuverability are severely restricted')
     
    	hdr_cells = table.rows[5].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('E')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('traffic flow is reduced to a slow but relatively uniform speeds, and traffic volume is equal to or nearly equal to capacity and maneuverability is extremely difficult')
     
    	hdr_cells = table.rows[6].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('F')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('volume greatly exceeds the capacity and lengthy delays occur')
    
    	p = document.add_paragraph()
    	p.style = "SpaceBefore"
    	p.add_run("""The project was evaluated for the potential to result in increased CO concentrations in the project area. Based on LOS estimates, it has been determined that this project would not increase traffic congestion, increase idle emissions, or CO concentrations. """)
    	p.add_run("""The estimated LOS under the No-Build Alternative would be D, while the estimated LOS under the Build Alternative would be C.""").font.color.rgb = red
    else:
    	document.add_heading('Introduction', 4)
    	p = document.add_paragraph()
    	p.add_run('The EPA first set air quality standards for CO in 1971. For protection of both public health and welfare, EPA set an eight-hour primary standard at 9 parts per million (ppm) and a one-hour primary standard at 35 ppm. Nationally and, particularly in urban areas, the majority of CO emissions to ambient air come from mobile sources.')
    	
    	p = document.add_paragraph()
    	p.add_run('Georgia is in attainment for CO; however, CO is also a concern in areas where signalized intersections (due to idling vehicles) are operating at a Level-of-Service (LOS) D, E, or F in the project design year (20 year design horizon).')
    	
    	p = document.add_paragraph()
    	p.add_run("""The LOS is a standard means of classifying traffic conditions associated with various traffic volume levels and traffic flow conditions. There are six levels of service at which a roadway can operate, represented by the letters "A" through "F." Each level is defined by a maximum value for the ratio of traffic volume (V) to facility capacity (C) (""")
    	p.add_run("""see Table 2. LOS, below""").font.color.rgb = red     
    	p.add_run("""). The LOS for signalized intersections is determined by calculating the average control delay per vehicle for the intersection, i.e., the average amount of time it takes a vehicle to get through the intersection.""")
    	
    	p = document.add_paragraph()
    	p.style = "FigureCaption"
    	p.add_run('Table 2. Level of Service (LOS)')
    	table = document.add_table(rows=7, cols=2)
    	table.style = 'LightShading'
    	table.alignment = WD_TABLE_ALIGNMENT.CENTER
    	hdr_cells = table.rows[0].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Level of Service')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Definition')
     
    	hdr_cells = table.rows[1].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('A')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('volume is well below capacity and traffic is flowing freely')
     
    	hdr_cells = table.rows[2].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('B')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('volume is steady, the presence of other vehicles begins to be noticeable')
     
    	hdr_cells = table.rows[3].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('C')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('steady traffic flow, speeds and maneuverability are more closely controlled by traffic volumes')
     
    	hdr_cells = table.rows[4].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('D')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('approaching an unsteady flow in which speed and maneuverability are severely restricted')
     
    	hdr_cells = table.rows[5].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('E')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('traffic flow is reduced to a slow but relatively uniform speeds, and traffic volume is equal to or nearly equal to capacity and maneuverability is extremely difficult')
     
    	hdr_cells = table.rows[6].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('F')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('volume greatly exceeds the capacity and lengthy delays occur')
    	
    	p = document.add_paragraph()
    	p.style = "SpaceBefore"
    	p.add_run("""The proposed project has %s signalized intersection(s) that would operate at a LOS of D, E, or F in the %s design year.""" % (strIntersection, strDesignYear))
    	
    	p = document.add_paragraph()
    	p.add_run("""The intersection of %s has the highest traffic volume at the worst level of service (""" % (strIntersection))
    	p.add_run("""see Table 3, below""").font.color.rgb = red
    	p.add_run("""). Therefore, this intersection has been chosen as a representative intersection. The results of this intersection are considered the worst case for the project corridor.""")

    	p = document.add_paragraph()
    	p.style = "FigureCaption"
    	p.add_run('Table 3. Signalized Intersections (LOS)')
    	table = document.add_table(rows=2, cols=3)
    	table.style = 'LightShading'
    	table.alignment = WD_TABLE_ALIGNMENT.CENTER
    	hdr_cells = table.rows[0].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Signalized Intersection')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Daily Traffic')
    	paragraph = hdr_cells[2].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('LOS')

    	hdr_cells = table.rows[1].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Roberts Drive at Main Street')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('ADT or AADT Traffic')
    	paragraph = hdr_cells[2].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('D, E, or F')

    	document.add_heading('Assessment Methodology', 4)
    	p = document.add_paragraph()
    	p.add_run('The CAL3QHC model combines the California Line Source Model (CALINE3) developed by the California Department of Transportation, with an algorithm for estimating queue lengths at signalized intersections. The CALINE3 and CAL3QHC models are accepted by the EPA and the Federal Highway Administration (FHWA) as techniques for assessing the air quality impacts that may occur from the operation of motor vehicles on roadways.')
    	
    	document.add_heading('CO Microscale Model Input Parameters', 4)
    	p = document.add_paragraph()
    	p.add_run("""Inputs to the models were such that they would provide a "worst case" analysis. The term "worst case" is frequently used in air quality impact analyses. The approach is to use a set of "worst case" meteorological conditions: lowest realistic wind speed, worst reasonable stability class, lowest reasonable temperature, highest expected traffic volumes, emissions associated with peak speeds, and closest reasonable receptor locations. If the "worst case" concentration does not violate air quality standards, it can be reasonably assumed that under any future set of actual meteorological conditions, the actual air quality will be better than the standards.""")
    	
    	#Add Appendix info
    	p = document.add_paragraph()
    	p.add_run("""The emission factors used in the CO microscale model were based off of GDOT Rate Tables using the EPA's Motor Vehicle Emission Simulator (MOVES). Emission levels were calculated per site specific criteria, including road grade (0 percent), vehicle mix (%s passenger cars; %s heavy trucks; %s medium trucks), design year (%s), road type (%s), and speed limit (%s). All emissions factors were based on temperatures under 70 degrees in the %s (""" % (strAuto, strHeavy, strMedium, strDesignYear, strRoadType, strSpeedLimitRange, strRegion))
    	p.add_run("""see Appendix B for emission factor worksheets).""").font.color.rgb = red 
    	
    	p = document.add_paragraph()
    	p.add_run("""Meteorological inputs to the CAL3QHC model were those that would give the worst case CO concentrations. The wind angle to the roadway was modeled at 10 degree intervals with a wind speed of one meter per second. In general, atmospheric stability is a function of the temperature distribution with height, solar radiation, cloud cover, and wind speed. Stability is identified by six classes ranging from A (very unstable) to F (very stable). A Stability Class %s was used for this project. Stable atmospheres contain little turbulence in which pollutant concentrations are high.""" % (strStabilityClass))
    	
    	p = document.add_paragraph()
    	p.add_run("""A mixing cell height (the elevation of the boundary between the vertically mixed layer of air closest to the earths surface and the relatively stable layer of air above) of 1,000 meters, the default value of CAL3QHC, and a surface roughness (the proportional measure of the height of obstacles to the wind flow) factor of %s cm was used (""" % (strSurfaceRoughness))
    	p.add_run("""see Table 4, below""").font.color.rgb = red
    	p.add_run("""). All roadway segments were modeled as at-grade facilities. Peak PM hourly traffic volumes were used as the worst case conditions for the one-hour analysis.""")
    	
    	#Add Appendix info
    	p = document.add_paragraph()
    	p.add_run("""The background concentration is usually defined as the concentration immediately upwind of the source. Through an agreement with the Georgia EPD, background CO concentration is considered to be only a small portion of the total input to the micro scale analysis. A background concentration of 1 ppm is added for rural areas, 2 ppm for suburban areas, and 3 ppm for urban areas. This area is SUBURBAN, therefore a background concentration of %s ppm has been added to the air quality dispersion modeling results. The intersection of %s was analyzed for the %s existing, %s build and no-build alternatives (""" % (strBackgroundConc, strIntersection, strExistingYear, strDesignYear))
    	p.add_run("""see Appendix C for CO inputs).""").font.color.rgb = red
    		
    	p = document.add_paragraph()
    	p.add_run('Receptor locations were identified to perform pollutant calculations. There were %s receptors chosen in the no-build conditions and %s receptors were chosen in the build condition. These locations were chosen because they represent the closest location where the public is likely to be present. Due to their proximity to the signal controlled intersection, these receptors would provide the highest concentrations of CO from the operations of the roadways.' % (strBuildRec, strNoBuildRec))
    	
    	p = document.add_paragraph()
    	p.style = "FigureCaption"
    	p.add_run('Table 4. Surface Roughness Lengths for Various Land Uses')
    	table = document.add_table(rows=6, cols=2)
    	table.style = 'LightShading'
    	table.alignment = WD_TABLE_ALIGNMENT.CENTER
    	hdr_cells = table.rows[0].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('City Land Use Surface Type')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Surface Roughness Length (cm)')
     
    	hdr_cells = table.rows[1].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Apartment Residential')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('370')
     
    	hdr_cells = table.rows[2].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Central Business District')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('321')
     
    	hdr_cells = table.rows[3].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Office')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('175')
     
    	hdr_cells = table.rows[4].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Park')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('127')
     
    	hdr_cells = table.rows[5].cells
    	paragraph = hdr_cells[0].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('Single Family Residential')
    	paragraph = hdr_cells[1].paragraphs[0]
    	paragraph.style = 'TblCentered'
    	run = paragraph.add_run('108')
    	
    	p = document.add_paragraph()
    	p.style = "SpaceBefore"
    	p.add_run('Sensitive receptors, defined as hospitals, nursing homes, schools, and other institutional facilities, were also considered as a part of this study. No facilities of these types are located proximate to the studied intersection; therefore, no special considerations were made for sensitive receptors.')
    	
    	#Add Appendix Data
    	document.add_heading('CO Microscale Model Summary of Results', 4)
    	p = document.add_paragraph()
    	p.add_run("""The State of Georgia and the EPA have set the maximum acceptable average CO concentrations at 35 ppm for a one-hour period, and 9 ppm for a continuous eight-hour period. The peak one-hour concentrations for the %s existing, and %s build and no-build conditions were predicted (""" % (strExistingYear, strDesignYear))
    	p.add_run("""see Table 5., below""").font.color.rgb = red
    	p.add_run("""). A copy of the model run data input and outputs are included in """)
    	p.add_run("""Appendix C""").font.color.rgb = red
    	p.add_run(""". The highest 1-hour CO concentration of %s in the %s design year is projected at receptor %s, located in the %s quadrant of the studied intersection. This value is lower than the maximum allowable NAAQS for the one-hour level of 35 ppm and the eight-hour level of 9 ppm. Since the highest one-hour concentration is lower than both the one-hour and eight-hour standards, an eight-hour concentration was not calculated. (Note eight-hour concentrations are calculated by multiplying the CAL3QHC results by a persistent factor of 0.6 and adding the background concentration to the results.)""" % (strCOBuild, strDesignYear, strMaxReceptorBuild, strReceptorQuadrant))
    	
    	'''Insert CO Graphics'''
    	insertcographics(document, strCOWorkspace)
    	
    	'''Insert CO Table'''
    	p = document.add_paragraph()
    	p.style = "FigureCaption"
    	p.add_run('Table 5. Predicted Highest One-Hour CO concentrations (ppm)')
    	createcotable(document, strCOWorkspace, int(strBackgroundConc))
    	
    	p = document.add_paragraph()
    	p.style = "FigureNote"
    	p.add_run("""Note: a background concentration of %s ppm has been added to the air quality model results.""" % (strBackgroundConc))
    	p = document.add_paragraph()
    	p.add_run('The results of the microscale analysis for the project area demonstrated that CO concentrations would not exceed state or federal air quality standards through the predicted design year traffic estimates. It is concluded that this project is consistent with region wide air quality goals and is consistent with the SIP for air quality.')
    	
    hd = document.add_heading('Is this project in a PM', 2)
    hd.add_run('2.5').font.subscript = True
    hd.add_run(' Non-Attainment Area?')
    if boolPM == False:
    	p = document.add_paragraph()
    	p.add_run('This project is located within a PM')
    	p.add_run('2.5').font.subscript = True
    	p.add_run(' attainment area. Qualitative PM')
    	p.add_run('2.5').font.subscript = True
    	p.add_run(' assessments are only required for projects of air quality concern within the PM')
    	p.add_run('2.5').font.subscript = True
    	p.add_run(' non-attainment area. Therefore, an assessment is not required.')
    else:
    	if strPMDet:
    		p = document.add_paragraph()
    		p.add_run('Transportation conformity is required for federal transportation projects in areas that have been designated by the EPA as not meeting the NAAQS. These areas are called non-attainment areas if they currently do not meet air quality standards or maintenance areas if they have previously violated air quality standards, but currently meet them and have an approved maintenance plan. On January 5, 2005, The EPA designated 24 counties and three partial counties in Georgia as non-attainment areas for fine particular matter, called PM')
    		p.add_run('2.5').font.subscript = True
    		p.add_run(""". This designation became effective on April 5, 2005, 90 days after EPA's published action in the Federal Register. Transportation Conformity for the PM""")
    		p.add_run('2.5').font.subscript = True
    		p.add_run(' standards applies as of April 5, 2006, after the one year grace period provided by the CAA. Metropolitan PM')
    		p.add_run('2.5').font.subscript = True
    		p.add_run(' non-attainment areas are now required to have a TIP and long range transportation plan (LRTP) that conforms to the PM')
    		p.add_run('2.5').font.subscript = True
    		p.add_run(' standard.')
    		#Add Appendix info

    		p = document.add_paragraph()
    		p.add_run("""This project has been evaluated by an interagency group consisting of FHWA, EPA, Georgia Department of Natural Resources Environmental Protection Division (GA EPD), and the Metropolitan Planning Organization (MPO) and they agreed that this project does NOT appear to be a "Project of Concern" per the Transportation Conformity Rule and thus meets the statutory and regulatory requirements for PM""")
    		p.add_run('2.5').font.subscript = True
    		p.add_run(""" hotspots without a quantitative analysis on %s. Documentation and correspondence are included in """ %(strConcDate))
    		p.add_run("""Appendix D""").font.color.rgb = red
    		p.add_run(""".""")
    		#Add Appendix info
    	else:
    		p = document.add_paragraph()
    		p.add_run('. This project has been evaluated by an interagency group consisting of FHWA, EPA, Georgia Department of Natural Resources Environmental Protection Division (GA EPD) and the MPO and was found to be exempt from the PM')
    		p.add_run('2.5').font.subscript = True
    		p.add_run(""" hot spot requirements on %s. Documentation and correspondence are included in """ % (strConcDate))
    		p.add_run("""Appendix D""").font.color.rgb = red
    		p.add_run(""".""")
    
    document.add_heading('How Does the Proposed Project Affect Mobile Source Air Toxics (MSAT)?', 2)    
    
    document.add_heading('Mobile Source Air Toxics (MSAT)', 3)
    
    if not boolMSAT:
    	p = document.add_paragraph()
    	p.add_run("""Mobile Source Air Toxics (MSAT) assessments are required statewide for most federal transportation projects. Based on the example projects defined in the FHWA guidance "Interim Guidance Update on Mobile Source Air Toxic Analysis in National Environmental Policy Act (NEPA) Documents" dated December 6, 2012, %s would be classified as a project with no meaningful MSAT impacts. In addition to the criteria air pollutants that must meet the NAAQS, EPA also regulates air toxics. Most air toxics originate from human-made sources, including on-road mobile sources, non-road mobile sources (e.g., airplanes), area sources (e.g., dry cleaners), and stationary sources (e.g., factories or refineries).""" % (strProjectName))
    	
    	p = document.add_paragraph()
    	p.add_run("""The purpose of this project is to %s. This project has been determined to generate minimal air quality impacts for CAA criteria pollutants and has not been linked with any special MSAT concerns. As such, this project will not result in changes in traffic volumes, vehicle mix, basic project location, or any other factor that would cause an increase in MSAT impacts of the project from that of the No-Build Alternative. It is therefore concluded the proposed action would have no meaningful potential MSAT impacts.""" % (strProjDesc))
    	
    	p = document.add_paragraph()
    	p.add_run("""Moreover, EPA regulations for vehicle engines and fuels will cause overall MSAT emissions to decline significantly over the next several decades. Based on regulations now in effect, an analysis of national trends with EPA's MOVES model forecasts a combined reduction of over 80 percent in the total annual emission rate for the priority MSAT from 2010 to 2050 while vehicle-miles of travel are projected to increase by over 100 percent. This will both reduce the background level of MSAT as well as the possibility of even minor MSAT emissions from this project.""")
    	
    else:
    	document.add_heading('Introduction', 4)	
    	p = document.add_paragraph()
    	p.add_run("""Mobile Source Air Toxics (MSAT) assessments are required statewide for most federal transportation projects. Based on the example projects defined in the FHWA guidance "Interim Guidance Update on Mobile Source Air Toxic Analysis in NEPA Documents" dated December 6, 2012, %s would be classified as a project with low potential MSAT effects. In addition to the criteria air pollutants that must meet the NAAQS, EPA also regulates air toxics. Most air toxics originate from human-made sources, including on-road mobile sources, non-road mobile sources (e.g., airplanes), area sources (e.g., dry cleaners) and stationary sources (e.g., factories or refineries).""" % (strProjectName))
    
    	document.add_heading('Background', 4)
    	p = document.add_paragraph()
    	p.add_run('Controlling air toxic emissions became a national priority with the passage of the Clean Air Act Amendments (CAAA) of 1990, whereby Congress mandated that the EPA regulate 188 air toxics, also known as hazardous air pollutants. The EPA has assessed this expansive list in their latest rule on the Control of Hazardous Air Pollutants from Mobile Sources (Federal Register, Vol. 72, No. 37, page 8430, February 26, 2007), and identified a group of 93 compounds emitted from mobile sources that are listed in their Integrated Risk Information System (IRIS) (http://www.epa.gov/iris/). In addition, EPA identified seven compounds with significant contributions from mobile sources that are among the national and regional-scale cancer risk drivers from their 1999 National Air Toxics Assessment (NATA) (http://www.epa.gov/ttn/atw/nata1999/). These are acrolein, benzene, 1,3-butadiene, diesel particulate matter plus diesel exhaust organic gases (diesel PM), formaldehyde, naphthalene, and polycyclic organic matter. While FHWA considers these the priority mobile source air toxics, the list is subject to change and may be adjusted in consideration of future EPA rules. The 2007 EPA rule mentioned above requires controls that will dramatically decrease MSAT emissions through cleaner fuels and cleaner engines.')
    
    	document.add_heading('Motor Vehicle Emissions Simulator (MOVES)', 4)
    	p = document.add_paragraph()
    	p.add_run("""According to EPA, MOVES improves upon the previous MOBILE model in several key aspects: MOVES is based on a vast amount of in-use vehicle data collected and analyzed since the latest release of MOBILE, including millions of emissions measurements from light-duty vehicles. Analysis of this data enhanced EPA's understanding of how mobile sources contribute to emissions inventories and the relative effectiveness of various control strategies. In addition, MOVES accounts for the significant effects that vehicle speed and temperature have on PM emissions estimates, whereas MOBILE did not. MOVES2010b includes all air toxic pollutants in NATA that are emitted by mobile sources. EPA has incorporated more recent data into MOVES2010b to update and enhance the quality of MSAT emission estimates. These data reflect advanced emission control technology and modern fuels, plus additional data for older technology vehicles.""")
    
    	p = document.add_paragraph()
    	p.add_run("""Based on an FHWA analysis using EPA's MOVES2010b model, as shown in Figure 2, even if vehicle miles traveled (VMT) increases by 102 percent as assumed from 2010 to 2050, a combined reduction of 83 percent in the total annual emissions for the priority MSAT is projected for the same time period.""")
    
    	p = document.add_paragraph()
    	p.style = "FigureCaption"
    	p.add_run('Figure 2. National MSAT Emission Trends 2010-2050 for Vehicles Operating on Roadways Using the EPA MOVES 2010b Model')
    	document.add_picture('../auxfiles/MOVES_Graphic.gif', width=docx.shared.Inches(4.00))
    	last_paragraph = document.paragraphs[-1] 
    	last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    	p = document.add_paragraph()
    	p.style = "FigureNote"
    	p.add_run('Note: Trends for specific locations may be different, depending on locally derived information representing vehicle-miles traveled, vehicle speeds, vehicle mix, fuels, emission control programs, meteorology, and other factors Source: EPA MOVES2010b model runs conducted during May - June 2012 by FHWA.')
    
    	p = document.add_paragraph()
    	p.add_run('The implications of MOVES on MSAT emissions estimates compared to MOBILE are: lower estimates of total MSAT emissions; significantly lower benzene emissions; significantly higher diesel PM emissions, especially for lower speeds. Consequently, diesel PM is projected to be the dominant component of the emissions total.')
    
    	p = document.add_paragraph()
    	p.add_run('Air toxics analysis is a continuing area of research. While much work has been done to assess the overall health risk of air toxics, many questions remain unanswered. In particular, the tools and techniques for assessing project-specific health outcomes as a result of lifetime MSAT exposure remain limited. These limitations impede the ability to evaluate how potential public health risks posed by MSAT exposure should be factored into project-level decision-making within the context of NEPA.')
    
    	p = document.add_paragraph()
    	p.add_run('Nonetheless, air toxics concerns continue to be raised on highway projects during the NEPA process. Even as the science emerges, we are duly expected by the public and other agencies to address MSAT impacts in our environmental documents. The FHWA, EPA, the Health Effects Institute, and others have funded and conducted research studies to try to more clearly define potential risks from MSAT emissions associated with highway projects. The FHWA will continue to monitor the developing research in this field.')
    
    	document.add_heading('Qualitative MSAT Assessment', 4)
    	p = document.add_paragraph()
    	p.add_run('For each alternative, the amount of MSAT emitted would be proportional to VMT, assuming that other variables such as fleet mix are the same for each alternative. The VMT estimated for the Build Alternatives is the same as that of the No-Build Alternative (')
    	p.add_run('see Table 6., below').font.color.rgb = red
    	p.add_run('). The emissions increase is offset somewhat by lower MSAT emission rates due to increased speeds; according to EPA\'s MOVES2010b model, emissions of all of the priority MSAT decrease as speed increases.')
					
    	p = document.add_paragraph()
    	p.style = "FigureCaption"
    	p.add_run('Table 6. VMT')
    	msattable(document, strRoadway, dblLength, intExistingADT, intNoBuildADT, intBuildADT)
    
    	p = document.add_paragraph()
    	p.style = "SpaceBefore"
    	p.add_run("""The estimated VMT under the Build Alternative is expected to be the same as that of the No-Build Alternative. It is expected there would be no appreciable difference in overall MSAT emissions among the two alternatives. Regardless of the alternative chosen, emissions will likely be lower than present levels in the Build year as a result of EPA's national control programs that are projected to reduce annual MSAT emissions by over 80 percent between 2010 and 2050. Local conditions may differ from these national projections in terms of fleet mix and turnover, VMT growth rates, and local control measures. However, the magnitude of the EPA projected reductions is so great (even after accounting for VMT growth) that MSAT emissions in the study area are likely to be lower in the future in nearly all cases.""")
    
    	p = document.add_paragraph()
    	p.add_run("""%s contemplated as part of the Build Alternative will have the effect of moving some traffic closer to nearby homes and businesses; therefore, under each alternative there may be localized areas where ambient concentrations of MSAT could be higher under certain Build Alternatives than the No-Build Alternative. The localized increases in MSAT concentrations would likely be most pronounced %s. However, the magnitude and the duration of these potential increases compared to the No-Build alternative cannot be reliably quantified due to incomplete or unavailable information in forecasting project-specific MSAT health impacts. In sum, when a highway is widened, the localized level of MSAT emissions for the Build Alternative could be higher relative to the No-Build Alternative, but this could be offset due to increases in speeds and reductions in congestion (which are associated with lower MSAT emissions). Also, MSAT will be lower in other locations when traffic shifts away from them. However, on a regional basis, EPA's vehicle and fuel regulations, coupled with fleet turnover, will over time cause substantial reductions that, in almost all cases, will cause region-wide MSAT levels to be significantly lower than today.""" % (strMSATDesc, strMSATConc))
    
    	document.add_heading('Incomplete or Unavailable Information for Project-Specific MSAT Health Impacts Analysis', 4)
    	p = document.add_paragraph()
    	p.add_run("""In FHWA's view, information is incomplete or unavailable to credibly predict the project-specific health impacts due to changes in MSAT emissions associated with a proposed set of highway alternatives. The outcome of such an assessment, adverse or not, would be influenced more by the uncertainty introduced into the process through assumption and speculation rather than any genuine insight into the actual health impacts directly attributable to MSAT exposure associated with a proposed action.""")
    
    	p = document.add_paragraph()
    	p.add_run("""The EPA is responsible for protecting the public health and welfare from any known or anticipated effect of an air pollutant. They are the lead authority for administering the CAA and its amendments and have specific statutory obligations with respect to hazardous air pollutants and MSAT. The EPA is in the continual process of assessing human health effects, exposures, and risks posed by air pollutants. They maintain the Integrated Risk Information System (IRIS), which is "a compilation of electronic reports on specific substances found in the environment and their potential to cause human health effects" (EPA, http://www.epa.gov/iris/). Each report contains assessments of non-cancerous and cancerous effects for individual compounds and quantitative estimates of risk levels from lifetime oral and inhalation exposures with uncertainty spanning perhaps an order of magnitude.""")
    
    	p = document.add_paragraph()
    	p.add_run("""Other organizations are also active in the research and analyses of the human health effects of MSAT, including the Health Effects Institute (HEI). Two HEI studies are summarized in Appendix D of FHWA's "Interim Guidance Update on Mobile source Air Toxic Analysis" in NEPA Documents. Among the adverse health effects linked to MSAT compounds at high exposures are; cancer in humans in occupational settings; cancer in animals; and irritation to the respiratory tract, including the exacerbation of asthma. Less obvious is the adverse human health effects of MSAT compounds at current environmental concentrations (HEI, http://pubs.healtheffects.org/view.php?id=282) or in the future as vehicle emissions substantially decrease (HEI, http://pubs.healtheffects.org/view.php?id=306).""")
    
    	p = document.add_paragraph()
    	p.add_run('The methodologies for forecasting health impacts include emissions modeling; dispersion modeling; exposure modeling; and then final determination of health impacts - each step in the process building on the model predictions obtained in the previous step. All are encumbered by technical shortcomings or uncertain science that prevents a more complete differentiation of the MSAT health impacts among a set of project alternatives. These difficulties are magnified for lifetime (i.e., 70 year) assessments, particularly because unsupportable assumptions would have to be made regarding changes in travel patterns and vehicle technology (which affects emissions rates) over that time frame, since such information is unavailable.')
    
    	p = document.add_paragraph()
    	p.add_run('It is particularly difficult to reliably forecast 70-year lifetime MSAT concentrations and exposure near roadways; to determine the portion of time that people are actually exposed at a specific location; and to establish the extent attributable to a proposed action, especially given that some of the information needed is unavailable.')
    
    	p = document.add_paragraph()
    	p.add_run('There are considerable uncertainties associated with the existing estimates of toxicity of the various MSAT, because of factors such as low-dose extrapolation and translation of occupational exposure data to the general population, a concern expressed by HEI (http://pubs.healtheffects.org/view.php?id=282). As a result, there is no national consensus on air dose-response values assumed to protect the public health and welfare for MSAT compounds, and in particular for diesel PM. The EPA (http://www.epa.gov/risk/ basicinformation.htm#g) and the HEI (http://pubs.healtheffects.org/getfile.php?u=395) have not established a basis for quantitative risk assessment of diesel PM in ambient settings.')
    
    	p = document.add_paragraph()
    	p.add_run("""There is also the lack of a national consensus on an acceptable level of risk. The current context is the process used by the EPA as provided by the CAA to determine whether more stringent controls are required in order to provide an ample margin of safety to protect public health or to prevent an adverse environmental effect for industrial sources subject to the maximum achievable control technology standards, such as benzene emissions from refineries. The decision framework is a two-step process. The first step requires EPA to determine an "acceptable" level of risk due to emissions from a source, which is generally no greater than approximately 100 in a million. Additional factors are considered in the second step, the goal of which is to maximize the number of people with risks less than one in a million due to emissions from a source. The results of this statutory two-step process do not guarantee that cancer risks from exposure to air toxics are less than one in a million; in some cases, the residual risk determination could result in maximum individual cancer risks that are as high as approximately 100 in a million. In a June 2008 decision, the U.S. Court of Appeals for the District of Columbia Circuit upheld EPA's approach to addressing risk in its two step decision framework. Information is incomplete or unavailable to establish that even the largest of highway projects would result in levels of risk greater than deemed acceptable.""")
    
    	p = document.add_paragraph()
    	p.add_run('Because of the limitations in the methodologies for forecasting health impacts described, any predicted difference in health impacts between alternatives is likely to be much smaller than the uncertainties associated with predicting the impacts. Consequently, the results of such assessments would not be useful to decision makers, who would need to weigh this information against project benefits, such as reducing traffic congestion, accident rates, and fatalities plus improved access for emergency response, that are better suited for quantitative analysis.')
    
    document.add_heading('How Does the Construction of this Project Affect Air Quality?', 4)
    p = document.add_paragraph()
    p.add_run('All phases of construction operations would temporarily contribute to air pollution. Particulates would increase slightly in the corridor as dust from construction collects in the air surrounding the project. The construction equipment would also produce slight amounts of exhaust emissions. The Rules and Regulations for Air Quality Control outlined in Chapter 391-3-1, Rules of GA EPD, would be followed during the construction of the project. These include covering earth-moving trucks to keep dust levels down, watering haul roads, and refraining from open burning, except as may be permitted by local regulations. ')
    
    p = document.add_paragraph()
    p.add_run('The EPA has listed a number of approved diesel retrofit technologies; many of these can be deployed as emissions mitigation measures for equipment used in construction. This listing can be found at: http://www.epa.gov/cleandiesel/technologies/retrofits.htm.')
    
    document.add_heading('What are the Conclusions Reached Based on the Air Assessment?', 4)
    p = document.add_paragraph()
    p.add_run('This project was evaluated for its consistency with state and federal air quality goals, including CO, Ozone, PM')
    p.add_run('2.5').font.subscript = True
    p.add_run(', and MSATs as part of this assessment. Results indicated that the project is consistent with the SIP for the attainment of clean air quality in Georgia and is in compliance with both state and federal air quality standards.')

if __name__ == "__main__":
    executivesummary(document)
    reportbody(document)
    document.save(os.path.join(localpath,'air_report.docx'))