# -*- coding: utf-8 -*-
"""
Created on Thu Jan 29 13:11:38 2015

@author: bbatt
"""
import os, docx
from docx.shared import Inches
from math import ceil

#existingout = os.path.join(workspace,"Existing.out")
#nobuildout = os.path.join(workspace,"NoBuild.out")
#buildout =  os.path.join(workspace,"Build.out")
#outputs = (existingout,nobuildout,buildout)
#document = docx.Document()
def maxreceptorcount(tploutputs):
    x = 0
    for item in tploutputs:
        if len(item) > x:
            x = len(item)
    return x

def parsecoresults(workspace, backgroundppm):
    existingout = os.path.join(workspace,"Existing.out")
    nobuildout = os.path.join(workspace,"NoBuild.out")
    buildout =  os.path.join(workspace,"Build.out")
    outputs = (existingout,nobuildout,buildout)    
    alloutputs = []
    for output in outputs:            
        f = open(output,"rb")
        with f:
            linelist = [i for i in f.readlines() if i.startswith(" MAX")]
            cooutputlist = [i for i in linelist[0].split() if i != '*' and i != "MAX"]
            convertedlist = [float(i) + backgroundppm for i in cooutputlist]
        alloutputs.append(convertedlist)
    return alloutputs

def insertcographics(doc, workspace):
    document = doc
    table = document.add_table(rows=1, cols=2)
    row_cells = table.rows[0].cells
    existingp = row_cells[0].paragraphs[0]
    existingp.style = 'TblCentered'
    existingr = existingp.add_run()
    existingr.add_picture(os.path.join(workspace,'existing.png'), width = Inches(3.32))
    existingp.add_run("Figure 2. Existing/NoBuild")
    buildp = row_cells[1].paragraphs[0]
    buildp.style = 'TblCentered'
    buildr = buildp.add_run()
    buildr.add_picture(os.path.join(workspace,'build.png'), width = Inches(3.32), height = Inches(3.32))
    buildp.add_run("Figure 3. Build")

def createcotable(doc, workspace, backgroundppm):
    document = doc
    outputsources = parsecoresults(workspace, backgroundppm)
    requiredtblrows = maxreceptorcount(outputsources)
    table = document.add_table(rows=1, cols=4)
    table.style = 'LightShading'
    hdrcells = table.rows[0].cells
    paragraph = hdrcells[0].paragraphs[0]
    paragraph.style = 'TblCentered'
    paragraph.add_run("Receptors")
    paragraph = hdrcells[1].paragraphs[0]
    paragraph.style = 'TblCentered'
    paragraph.add_run("Existing")
    paragraph = hdrcells[2].paragraphs[0]
    paragraph.style = 'TblCentered'
    paragraph.add_run("No Build")
    paragraph = hdrcells[3].paragraphs[0]
    paragraph.style = 'TblCentered'
    paragraph.add_run("Build")
    for i in range(requiredtblrows):
        row_cells = table.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        paragraph.style = 'TblCentered'
        paragraph.add_run("Receptor " + str(i+1))
        paragraph = row_cells[1].paragraphs[0]
        paragraph.style = 'TblCentered'
        paragraph.add_run(str(outputsources[0][i]))
        paragraph = row_cells[2].paragraphs[0]
        paragraph.style = 'TblCentered'
        paragraph.add_run(str(outputsources[1][i]))
        paragraph = row_cells[3].paragraphs[0]
        paragraph.style = 'TblCentered'
        paragraph.add_run(str(outputsources[2][i]))
    return table

def msattable(document, road, length, existing, nobuild, build):
	table = document.add_table(rows=3, cols=6)
	table.style = 'LightShading'
	hdrs = ["Roadway", "Roadway Length (Miles)", "", "Existing", "No Build", "Build" ]
	for i in range(len(hdrs)):
		table.rows[0].cells[i].text = hdrs[i]
	rowindex = 1
	roadcell = table.cell(rowindex,0).merge(table.cell(2,0))
	lengthcell = table.cell(rowindex, 1).merge(table.cell(2, 1))
	trafficlabelcell = table.cell(rowindex, 2)	
	existingcell = table.cell(rowindex,3)
	nobuildcell = table.cell(rowindex,4)
	buildcell = table.cell(rowindex,5)
	vmtlabelcell = table.cell(rowindex+ 1, 2)	
	existingvmtcell = table.cell(rowindex + 1, 3)
	nobuildvmtcell = table.cell(rowindex + 1, 4)
	buildvmtcell = table.cell(rowindex + 1, 5)
	roadcell.text = road
	lengthcell.text = str(length)
	trafficlabelcell.text = "ADT"
	existingcell.text = "{:,}".format(existing)
	nobuildcell.text = "{:,}".format(nobuild)
	buildcell.text = "{:,}".format(build)
	vmtlabelcell.text = "VMT"
	existingvmtcell.text = "{:,}".format(ceil(existing * length)).rstrip(".").rstrip("0").rstrip(".")
	nobuildvmtcell.text = "{:,}".format(ceil(nobuild * length)).rstrip(".").rstrip("0").rstrip(".")
	buildvmtcell.text = "{:,}".format(ceil(build * length)).rstrip(".").rstrip("0").rstrip(".")
	return table

		